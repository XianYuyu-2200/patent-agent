import inspect
import json
import os
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from typer.testing import CliRunner

from codex_patent.cli import app
from codex_patent.export_docx import SECTION_ORDER, TEMPLATE_PATH, export_application
from codex_patent.models import ArtifactRef, PatentCase, ReviewIssue
from codex_patent.repository import CaseRepository


SECTIONS = {
    "技术领域": "本申请涉及支架技术领域。",
    "背景技术": "现有支架存在稳定性不足的问题。",
    "发明内容": "本申请提供一种折叠支架。",
    "附图说明": "图1为折叠支架的结构示意图。",
    "具体实施方式": "如图1所示，支架包括底座和支撑件。",
    "摘要": "本申请公开一种折叠支架。",
}
CLAIMS_TEXT = "\n\n".join(
    [
        "1. 一种折叠支架，其特征在于，包括底座和支撑件。",
        "2. 根据权利要求1所述的折叠支架，其特征在于，所述支撑件可转动。",
    ]
)
SPECIFICATION_TEXT = "\n\n".join(
    value
    for heading in SECTION_ORDER[:-1]
    for value in (heading, SECTIONS[heading])
)
ARTIFACT_PATHS = {
    "claims": "drafts/claims-v2.md",
    "specification": "drafts/specification-v2.md",
    "abstract": "drafts/abstract-v2.md",
    "quality-review": "review-log/quality-review-v2.json",
}
ARTIFACT_CONTENT = {
    "claims": CLAIMS_TEXT,
    "specification": SPECIFICATION_TEXT,
    "abstract": SECTIONS["摘要"],
}
RUNNER = CliRunner()


def _official_quality_review(version: int = 2, **overrides: object) -> dict:
    review = {
        "review_status": "completed",
        "input_versions": {
            "claims": f"v{version}",
            "claim_feature_map": f"v{version}",
            "specification": f"v{version}",
            "abstract": f"v{version}",
            "drawing_plan": f"v{version}",
            "prior_art": f"v{version}",
        },
        "checks": {
            check: {
                "status": "completed",
                "conclusion_or_gap": "reviewed with no delivery-blocking finding",
                "source_anchors": [f"{check}-v{version}"],
            }
            for check in (
                "claim_clarity_dependency",
                "specification_support_sufficiency",
                "cross_artifact_consistency",
                "abstract_fidelity",
                "unity",
                "subject_matter_technical_solution",
                "filing_type",
                "novelty",
                "inventive_step",
                "design_around",
                "form",
            )
        },
        "findings": [],
        "open_issue_counts": {
            "critical": 0,
            "high": 0,
            "medium": 0,
            "low": 0,
        },
        "delivery_recommendation": "ready-for-human-review",
        "unresolved_questions": [],
        "source_anchors": [
            f"claims-v{version}.md",
            f"specification-v{version}.md",
            f"abstract-v{version}.md",
        ],
    }
    review.update(overrides)
    return review


def _scoped_final_delivery_approval(version: int = 2) -> dict:
    return {
        "approval_id": f"FD-TEST-v{version}",
        "type": "final-delivery",
        "status": "approved",
        "current": True,
        "application_set": f"CN-TEST-001-v{version}",
        "scope": {
            "claims": f"v{version}",
            "specification": f"v{version}",
            "abstract": f"v{version}",
            "quality_review": f"v{version}",
            "action": "DOCX export",
        },
    }


def _replace_persisted_approvals(case_dir: Path, approvals: list[dict]) -> None:
    case_path = case_dir / "case.json"
    data = json.loads(case_path.read_text(encoding="utf-8"))
    data["approvals"] = approvals
    case_path.write_text(
        json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8"
    )


def _write_ready_case(
    tmp_path: Path,
) -> tuple[CaseRepository, PatentCase, Path]:
    repository = CaseRepository(tmp_path / "cases")
    case = PatentCase(
        case_id="CN-TEST-001",
        title="折叠支架",
        approvals=[_scoped_final_delivery_approval(version=2)],
        artifacts=[
            ArtifactRef(
                artifact_type=artifact_type,
                version=2,
                path=path,
            )
            for artifact_type, path in ARTIFACT_PATHS.items()
        ],
    )
    case_dir = repository.create(case)
    for artifact_type, relative_path in ARTIFACT_PATHS.items():
        artifact_path = case_dir / relative_path
        artifact_path.parent.mkdir(parents=True, exist_ok=True)
        content = (
            json.dumps(_official_quality_review(), ensure_ascii=False)
            if artifact_type == "quality-review"
            else ARTIFACT_CONTENT[artifact_type]
        )
        artifact_path.write_text(content, encoding="utf-8")
    return repository, case, case_dir


def _artifact(case: PatentCase, artifact_type: str) -> ArtifactRef:
    return next(
        artifact
        for artifact in case.artifacts
        if artifact.artifact_type == artifact_type
    )


def _write_application_section(
    case: PatentCase,
    case_dir: Path,
    heading: str,
    content: str,
) -> None:
    sections = dict(SECTIONS)
    sections[heading] = content
    if heading == "摘要":
        abstract = case_dir / _artifact(case, "abstract").path
        abstract.write_text(content, encoding="utf-8")
        return

    specification = case_dir / _artifact(case, "specification").path
    specification.write_text(
        "\n\n".join(
            value
            for section_heading in SECTION_ORDER[:-1]
            for value in (section_heading, sections[section_heading])
        ),
        encoding="utf-8",
    )


def _assert_protected_output_is_refused_and_preserved(
    case_dir: Path,
    protected_path: Path,
    *,
    template_path: Path | None = None,
) -> None:
    original_bytes = protected_path.read_bytes()
    try:
        with pytest.raises(ValueError, match="protected"):
            export_application(
                case_dir,
                protected_path,
                final_approval=True,
                template_path=template_path,
            )
    finally:
        assert protected_path.read_bytes() == original_bytes


def _export_kwargs(tmp_path: Path) -> dict:
    _, _, case_dir = _write_ready_case(tmp_path)
    return {
        "case_dir": case_dir,
        "output_path": tmp_path / "application.docx",
        "final_approval": True,
    }


def test_export_contains_required_sections_from_persisted_case_workspace(
    tmp_path: Path,
):
    kwargs = _export_kwargs(tmp_path)
    output = export_application(**kwargs)

    text = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)
    for heading in (
        "权利要求书",
        "技术领域",
        "背景技术",
        "发明内容",
        "附图说明",
        "具体实施方式",
        "摘要",
    ):
        assert heading in text

    paragraphs = Document(output).paragraphs
    observed = [paragraph.text for paragraph in paragraphs]
    assert observed.index("权利要求书") < observed.index("技术领域")
    for previous, following in zip(SECTION_ORDER, SECTION_ORDER[1:]):
        assert observed.index(previous) < observed.index(following)
    assert CLAIMS_TEXT.split("\n\n") == [
        paragraph.text
        for paragraph in paragraphs
        if paragraph.style.name == "Patent Claim"
    ]


def test_export_refuses_absent_explicit_final_approval_before_opening_template(
    tmp_path: Path,
):
    kwargs = _export_kwargs(tmp_path)
    kwargs["final_approval"] = False
    kwargs["template_path"] = tmp_path / "missing-template.docx"

    with pytest.raises(ValueError, match="explicit final approval"):
        export_application(**kwargs)


def test_export_refuses_non_boolean_final_approval_before_opening_template(
    tmp_path: Path,
):
    kwargs = _export_kwargs(tmp_path)
    kwargs["final_approval"] = "yes"
    kwargs["template_path"] = tmp_path / "missing-template.docx"

    with pytest.raises(ValueError, match="actual boolean"):
        export_application(**kwargs)


def test_export_refuses_missing_persisted_final_delivery_approval(
    tmp_path: Path,
):
    repository, case, case_dir = _write_ready_case(tmp_path)
    case.approvals.clear()
    repository.save(case)

    with pytest.raises(ValueError, match="persisted final-delivery approval"):
        export_application(
            case_dir,
            tmp_path / "application.docx",
            final_approval=True,
            template_path=tmp_path / "missing-template.docx",
        )


def test_export_refuses_open_high_issue_from_persisted_case_state(
    tmp_path: Path,
):
    repository, case, case_dir = _write_ready_case(tmp_path)
    case.issues.append(
        ReviewIssue(
            issue_id="support-gap",
            severity="high",
            message="claim lacks support",
        )
    )
    repository.save(case)

    with pytest.raises(ValueError, match="open high-severity"):
        export_application(
            case_dir,
            tmp_path / "application.docx",
            final_approval=True,
            template_path=tmp_path / "missing-template.docx",
        )


@pytest.mark.parametrize("artifact_type", ARTIFACT_PATHS)
def test_export_refuses_missing_required_artifact_reference(
    tmp_path: Path,
    artifact_type: str,
):
    repository, case, case_dir = _write_ready_case(tmp_path)
    case.artifacts = [
        artifact
        for artifact in case.artifacts
        if artifact.artifact_type != artifact_type
    ]
    repository.save(case)

    with pytest.raises(ValueError, match=f"missing required artifact: {artifact_type}"):
        export_application(
            case_dir,
            tmp_path / "application.docx",
            final_approval=True,
            template_path=tmp_path / "missing-template.docx",
        )


@pytest.mark.parametrize("artifact_type", ARTIFACT_PATHS)
def test_export_refuses_stale_required_artifact(
    tmp_path: Path,
    artifact_type: str,
):
    repository, case, case_dir = _write_ready_case(tmp_path)
    _artifact(case, artifact_type).stale = True
    repository.save(case)

    with pytest.raises(ValueError, match=f"stale {artifact_type} artifact"):
        export_application(
            case_dir,
            tmp_path / "application.docx",
            final_approval=True,
            template_path=tmp_path / "missing-template.docx",
        )


def test_export_preserves_stale_docx_delivery_guard(tmp_path: Path):
    repository, case, case_dir = _write_ready_case(tmp_path)
    case.artifacts.append(
        ArtifactRef(
            artifact_type="docx",
            version=1,
            path="exports/application-v1.docx",
            stale=True,
        )
    )
    repository.save(case)

    with pytest.raises(ValueError, match="stale docx artifact"):
        export_application(
            case_dir,
            tmp_path / "application.docx",
            final_approval=True,
            template_path=tmp_path / "missing-template.docx",
        )


@pytest.mark.parametrize("artifact_type", ARTIFACT_PATHS)
def test_export_refuses_missing_referenced_artifact_file(
    tmp_path: Path,
    artifact_type: str,
):
    _, case, case_dir = _write_ready_case(tmp_path)
    (case_dir / _artifact(case, artifact_type).path).unlink()

    with pytest.raises(ValueError, match=f"missing {artifact_type} artifact file"):
        export_application(
            case_dir,
            tmp_path / "application.docx",
            final_approval=True,
            template_path=tmp_path / "missing-template.docx",
        )


@pytest.mark.parametrize("artifact_type", ARTIFACT_PATHS)
def test_export_refuses_empty_referenced_artifact_file(
    tmp_path: Path,
    artifact_type: str,
):
    _, case, case_dir = _write_ready_case(tmp_path)
    (case_dir / _artifact(case, artifact_type).path).write_text(
        "", encoding="utf-8"
    )

    with pytest.raises(ValueError, match=f"empty {artifact_type} artifact"):
        export_application(
            case_dir,
            tmp_path / "application.docx",
            final_approval=True,
            template_path=tmp_path / "missing-template.docx",
        )


def test_export_refuses_multiple_current_artifacts_of_same_type(tmp_path: Path):
    repository, case, case_dir = _write_ready_case(tmp_path)
    duplicate = ArtifactRef(
        artifact_type="claims",
        version=3,
        path="drafts/claims-v3.md",
    )
    case.artifacts.append(duplicate)
    (case_dir / duplicate.path).write_text(CLAIMS_TEXT, encoding="utf-8")
    repository.save(case)

    with pytest.raises(ValueError, match="multiple current claims artifacts"):
        export_application(case_dir, tmp_path / "application.docx", final_approval=True)


def test_export_refuses_version_mismatched_required_artifacts(tmp_path: Path):
    repository, case, case_dir = _write_ready_case(tmp_path)
    abstract = _artifact(case, "abstract")
    (case_dir / abstract.path).unlink()
    abstract.version = 3
    abstract.path = "drafts/abstract-v3.md"
    (case_dir / abstract.path).write_text(SECTIONS["摘要"], encoding="utf-8")
    repository.save(case)

    with pytest.raises(ValueError, match="required artifacts must share one version"):
        export_application(case_dir, tmp_path / "application.docx", final_approval=True)


def test_export_refuses_artifact_path_outside_case_workspace(tmp_path: Path):
    repository, case, case_dir = _write_ready_case(tmp_path)
    outside = case_dir.parent / "claims-v2.md"
    outside.write_text(CLAIMS_TEXT, encoding="utf-8")
    _artifact(case, "claims").path = "../claims-v2.md"
    repository.save(case)

    with pytest.raises(ValueError, match="inside the case workspace"):
        export_application(case_dir, tmp_path / "application.docx", final_approval=True)


def test_export_refuses_artifact_path_that_does_not_match_type_and_version(
    tmp_path: Path,
):
    repository, case, case_dir = _write_ready_case(tmp_path)
    _artifact(case, "claims").path = ARTIFACT_PATHS["specification"]
    repository.save(case)

    with pytest.raises(ValueError, match="claims artifact path must name claims-v2.md"):
        export_application(case_dir, tmp_path / "application.docx", final_approval=True)


def test_export_refuses_case_id_that_does_not_match_case_directory(tmp_path: Path):
    _, case, case_dir = _write_ready_case(tmp_path)
    case.case_id = "OTHER-CASE"
    (case_dir / "case.json").write_text(case.model_dump_json(indent=2), encoding="utf-8")

    with pytest.raises(ValueError, match="case ID does not match"):
        export_application(case_dir, tmp_path / "application.docx", final_approval=True)


def test_export_refuses_invalid_quality_review_json(tmp_path: Path):
    _, case, case_dir = _write_ready_case(tmp_path)
    quality_review = case_dir / _artifact(case, "quality-review").path
    quality_review.write_text("not-json", encoding="utf-8")

    with pytest.raises(ValueError, match="valid JSON object"):
        export_application(case_dir, tmp_path / "application.docx", final_approval=True)


def test_export_accepts_official_completed_quality_review_recipe(tmp_path: Path):
    _, case, case_dir = _write_ready_case(tmp_path)
    quality_review = case_dir / _artifact(case, "quality-review").path
    quality_review.write_text(
        json.dumps(_official_quality_review(), ensure_ascii=False),
        encoding="utf-8",
    )

    output = tmp_path / "application.docx"
    assert export_application(case_dir, output, final_approval=True) == output
    assert output.is_file()


@pytest.mark.parametrize(
    "artifact_type",
    ("claims", "specification", "abstract", "quality-review"),
)
def test_export_refuses_output_collision_with_current_artifact_and_preserves_bytes(
    tmp_path: Path,
    artifact_type: str,
):
    _, case, case_dir = _write_ready_case(tmp_path)
    protected_path = case_dir / _artifact(case, artifact_type).path

    _assert_protected_output_is_refused_and_preserved(case_dir, protected_path)


def test_export_refuses_output_collision_with_case_json_and_preserves_bytes(
    tmp_path: Path,
):
    _, _, case_dir = _write_ready_case(tmp_path)

    _assert_protected_output_is_refused_and_preserved(
        case_dir, case_dir / "case.json"
    )


def test_export_refuses_resolved_output_alias_and_preserves_bytes(tmp_path: Path):
    _, case, case_dir = _write_ready_case(tmp_path)
    claims = case_dir / _artifact(case, "claims").path
    alias = claims.parent / ".." / claims.parent.name / claims.name

    _assert_protected_output_is_refused_and_preserved(case_dir, alias)


@pytest.mark.skipif(os.name != "nt", reason="case-insensitive alias is Windows-specific")
def test_export_refuses_case_insensitive_output_alias_and_preserves_bytes(
    tmp_path: Path,
):
    _, case, case_dir = _write_ready_case(tmp_path)
    claims = case_dir / _artifact(case, "claims").path
    alias = claims.with_name(claims.name.upper())

    _assert_protected_output_is_refused_and_preserved(case_dir, alias)


def test_export_refuses_output_collision_with_source_material_and_preserves_bytes(
    tmp_path: Path,
):
    _, _, case_dir = _write_ready_case(tmp_path)
    source = case_dir / "source-materials" / "customer-source.docx"
    source.parent.mkdir(parents=True, exist_ok=True)
    source.write_bytes(b"customer source bytes must remain immutable")

    _assert_protected_output_is_refused_and_preserved(case_dir, source)


def test_export_refuses_output_collision_with_active_template_and_preserves_bytes(
    tmp_path: Path,
):
    _, _, case_dir = _write_ready_case(tmp_path)
    template = tmp_path / "active-template.docx"
    template.write_bytes(TEMPLATE_PATH.read_bytes())

    _assert_protected_output_is_refused_and_preserved(
        case_dir,
        template,
        template_path=template,
    )


def test_export_refuses_arbitrary_issues_only_quality_review_schema(tmp_path: Path):
    _, case, case_dir = _write_ready_case(tmp_path)
    quality_review = case_dir / _artifact(case, "quality-review").path
    quality_review.write_text(json.dumps({"issues": []}), encoding="utf-8")
    output = tmp_path / "application.docx"

    with pytest.raises(ValueError, match="official completed output recipe"):
        export_application(
            case_dir,
            output,
            final_approval=True,
            template_path=tmp_path / "missing-template.docx",
        )

    assert not output.exists()


def test_export_refuses_conflicting_legacy_issues_and_blocked_delivery_state(
    tmp_path: Path,
):
    _, case, case_dir = _write_ready_case(tmp_path)
    quality_review = case_dir / _artifact(case, "quality-review").path
    quality_review.write_text(
        json.dumps(
            {
                "issues": [],
                "review_status": "completed-with-issues",
                "delivery_recommendation": "blocked",
                "open_high": 2,
            }
        ),
        encoding="utf-8",
    )
    output = tmp_path / "application.docx"

    with pytest.raises(ValueError, match="official completed output recipe"):
        export_application(
            case_dir,
            output,
            final_approval=True,
            template_path=tmp_path / "missing-template.docx",
        )

    assert not output.exists()


def test_export_refuses_official_review_with_open_high_finding(tmp_path: Path):
    _, case, case_dir = _write_ready_case(tmp_path)
    quality_review = case_dir / _artifact(case, "quality-review").path
    review = _official_quality_review(
        review_status="completed-with-issues",
        findings=[
            {
                "issue_id": "QCR-SUP-001",
                "severity": "high",
                "category": "specification-support",
                "artifact": "specification-v2.md",
                "location": "claim 1",
                "source_anchors": ["specification-v2.md"],
                "explanation": "claim feature lacks support",
                "suggested_action": "human review",
                "blocks_delivery": True,
            }
        ],
        open_issue_counts={"critical": 0, "high": 1, "medium": 0, "low": 0},
        delivery_recommendation="blocked",
    )
    quality_review.write_text(json.dumps(review), encoding="utf-8")
    output = tmp_path / "application.docx"

    with pytest.raises(ValueError, match="not eligible for export"):
        export_application(
            case_dir,
            output,
            final_approval=True,
            template_path=tmp_path / "missing-template.docx",
        )

    assert not output.exists()


@pytest.mark.parametrize(
    "check_name",
    ("specification_support_sufficiency", "novelty"),
)
def test_export_refuses_delivery_check_that_is_not_assessable(
    tmp_path: Path,
    check_name: str,
):
    _, case, case_dir = _write_ready_case(tmp_path)
    quality_review = case_dir / _artifact(case, "quality-review").path
    review = _official_quality_review()
    review["checks"][check_name] = {
        "status": "not-assessable",
        "conclusion_or_gap": "required evidence is unavailable",
        "source_anchors": [{"evidence_gap": "required evidence is unavailable"}],
    }
    quality_review.write_text(json.dumps(review), encoding="utf-8")
    output = tmp_path / "application.docx"

    with pytest.raises(ValueError, match="not eligible for export"):
        export_application(
            case_dir,
            output,
            final_approval=True,
            template_path=tmp_path / "missing-template.docx",
        )

    assert not output.exists()


def test_export_refuses_blocked_input_integrity_metadata(tmp_path: Path):
    _, case, case_dir = _write_ready_case(tmp_path)
    quality_review = case_dir / _artifact(case, "quality-review").path
    review = _official_quality_review(
        input_integrity={
            "status": "blocked",
            "all_six_readable": False,
            "freshness": "stale",
            "mutually_version_matched": False,
            "internal_identity": "unidentifiable",
            "placeholder_block": True,
        }
    )
    quality_review.write_text(json.dumps(review), encoding="utf-8")
    output = tmp_path / "application.docx"

    with pytest.raises(ValueError, match="official completed output recipe"):
        export_application(
            case_dir,
            output,
            final_approval=True,
            template_path=tmp_path / "missing-template.docx",
        )

    assert not output.exists()


def test_export_refuses_empty_structured_source_anchor(tmp_path: Path):
    _, case, case_dir = _write_ready_case(tmp_path)
    quality_review = case_dir / _artifact(case, "quality-review").path
    review = _official_quality_review(source_anchors=[{}])
    quality_review.write_text(json.dumps(review), encoding="utf-8")
    output = tmp_path / "application.docx"

    with pytest.raises(ValueError, match="official completed output recipe"):
        export_application(
            case_dir,
            output,
            final_approval=True,
            template_path=tmp_path / "missing-template.docx",
        )

    assert not output.exists()


@pytest.mark.parametrize("anchor_location", ("top-level", "novelty-check"))
def test_export_refuses_blank_overlap_source_anchor(
    tmp_path: Path,
    anchor_location: str,
):
    _, case, case_dir = _write_ready_case(tmp_path)
    quality_review = case_dir / _artifact(case, "quality-review").path
    review = _official_quality_review()
    if anchor_location == "top-level":
        review["source_anchors"] = [{"overlap": [""]}]
    else:
        review["checks"]["novelty"]["source_anchors"] = [
            {"overlap": [""]}
        ]
    quality_review.write_text(json.dumps(review), encoding="utf-8")
    output = tmp_path / "application.docx"

    with pytest.raises(ValueError, match="official completed output recipe"):
        export_application(
            case_dir,
            output,
            final_approval=True,
            template_path=tmp_path / "missing-template.docx",
        )

    assert not output.exists()


def test_export_refuses_structurally_empty_verified_disclosure(tmp_path: Path):
    _, case, case_dir = _write_ready_case(tmp_path)
    quality_review = case_dir / _artifact(case, "quality-review").path
    review = _official_quality_review(
        prior_art_assessment={
            "verified_disclosures": [{}],
            "novelty_risk": [
                {"claim": "C1", "level": "none", "basis": "no overlap"}
            ],
            "inventive_step_risk": [
                {"claim": "C1", "level": "none", "basis": "no overlap"}
            ],
        }
    )
    quality_review.write_text(json.dumps(review), encoding="utf-8")

    with pytest.raises(ValueError, match="official completed output recipe"):
        export_application(
            case_dir,
            tmp_path / "application.docx",
            final_approval=True,
            template_path=tmp_path / "missing-template.docx",
        )


def test_export_refuses_completed_prior_art_checks_without_verified_disclosure(
    tmp_path: Path,
):
    _, case, case_dir = _write_ready_case(tmp_path)
    quality_review = case_dir / _artifact(case, "quality-review").path
    review = _official_quality_review(
        prior_art_assessment={
            "verified_disclosures": [],
            "novelty_risk": [
                {"claim": "C1", "level": "none", "basis": "none found"}
            ],
            "inventive_step_risk": [
                {"claim": "C1", "level": "none", "basis": "none found"}
            ],
        }
    )
    quality_review.write_text(json.dumps(review), encoding="utf-8")

    with pytest.raises(ValueError, match="official completed output recipe"):
        export_application(
            case_dir,
            tmp_path / "application.docx",
            final_approval=True,
            template_path=tmp_path / "missing-template.docx",
        )


@pytest.mark.parametrize(
    ("check_name", "risk_name"),
    (("novelty", "novelty_risk"), ("inventive_step", "inventive_step_risk")),
)
def test_export_refuses_completed_check_for_not_assessable_prior_art_risk(
    tmp_path: Path,
    check_name: str,
    risk_name: str,
):
    _, case, case_dir = _write_ready_case(tmp_path)
    quality_review = case_dir / _artifact(case, "quality-review").path
    review = _official_quality_review(
        prior_art_assessment={
            "verified_disclosures": [
                {
                    "document_id": "CN123456789A",
                    "disclosure_anchor": "[0042]-[0045]",
                    "disclosure": "verified disclosure",
                    "overlapping_features": ["F1"],
                }
            ],
            "novelty_risk": [
                {"claim": "C1", "level": "low", "basis": "limited overlap"}
            ],
            "inventive_step_risk": [
                {"claim": "C1", "level": "low", "basis": "limited overlap"}
            ],
        }
    )
    review["prior_art_assessment"][risk_name] = [
        {"claim": "C1", "level": "not-assessable", "gap": "missing evidence"}
    ]
    review["checks"][check_name] = {
        "status": "completed",
        "conclusion_or_gap": "missing evidence was treated as completed",
        "source_anchors": [{"evidence_gap": "missing evidence"}],
    }
    quality_review.write_text(json.dumps(review), encoding="utf-8")

    with pytest.raises(ValueError, match="official completed output recipe"):
        export_application(
            case_dir,
            tmp_path / "application.docx",
            final_approval=True,
            template_path=tmp_path / "missing-template.docx",
        )


def test_export_refuses_completed_prior_art_checks_with_only_evidence_gaps(
    tmp_path: Path,
):
    _, case, case_dir = _write_ready_case(tmp_path)
    quality_review = case_dir / _artifact(case, "quality-review").path
    review = _official_quality_review(
        prior_art_assessment={
            "verified_disclosures": [],
            "novelty_risk": [
                {"claim": "C1", "level": "not-assessable", "gap": "no evidence"}
            ],
            "inventive_step_risk": [
                {"claim": "C1", "level": "not-assessable", "gap": "no evidence"}
            ],
        }
    )
    for check_name in ("novelty", "inventive_step"):
        review["checks"][check_name] = {
            "status": "completed",
            "conclusion_or_gap": "no evidence was supplied",
            "source_anchors": [{"evidence_gap": "no evidence was supplied"}],
        }
    quality_review.write_text(json.dumps(review), encoding="utf-8")

    with pytest.raises(ValueError, match="official completed output recipe"):
        export_application(
            case_dir,
            tmp_path / "application.docx",
            final_approval=True,
            template_path=tmp_path / "missing-template.docx",
        )


def test_export_accepts_consistent_verified_prior_art_assessment(tmp_path: Path):
    _, case, case_dir = _write_ready_case(tmp_path)
    quality_review = case_dir / _artifact(case, "quality-review").path
    disclosure_anchor = {
        "artifact": "prior-art-v2.json",
        "document_id": "CN123456789A",
        "disclosure_anchor": "[0042]-[0045]",
        "overlap": ["F1"],
    }
    review = _official_quality_review(
        prior_art_assessment={
            "verified_disclosures": [
                {
                    "document_id": "CN123456789A",
                    "disclosure_anchor": "[0042]-[0045]",
                    "disclosure": "verified disclosure",
                    "overlapping_features": ["F1"],
                }
            ],
            "novelty_risk": [
                {"claim": "C1", "level": "low", "basis": "limited overlap"}
            ],
            "inventive_step_risk": [
                {"claim": "C1", "level": "low", "basis": "limited overlap"}
            ],
        }
    )
    for check_name in ("novelty", "inventive_step"):
        review["checks"][check_name]["source_anchors"] = [disclosure_anchor]
    quality_review.write_text(json.dumps(review), encoding="utf-8")
    output = tmp_path / "application.docx"

    assert export_application(case_dir, output, final_approval=True) == output
    assert output.is_file()


PLACEHOLDER_TEXTS = (
    "TODO",
    "placeholder",
    "...",
    "……",
    "待确认",
    "[待补充权利要求]",
)


@pytest.mark.parametrize("placeholder", PLACEHOLDER_TEXTS)
def test_export_refuses_placeholder_claim_text(
    tmp_path: Path,
    placeholder: str,
):
    _, case, case_dir = _write_ready_case(tmp_path)
    claims = case_dir / _artifact(case, "claims").path
    claims.write_text(placeholder, encoding="utf-8")
    output = tmp_path / "application.docx"

    with pytest.raises(ValueError, match="placeholder claim"):
        export_application(
            case_dir,
            output,
            final_approval=True,
            template_path=tmp_path / "missing-template.docx",
        )

    assert not output.exists()


@pytest.mark.parametrize("heading", SECTION_ORDER)
@pytest.mark.parametrize("placeholder", PLACEHOLDER_TEXTS)
def test_export_refuses_placeholder_required_section_text(
    tmp_path: Path,
    heading: str,
    placeholder: str,
):
    _, case, case_dir = _write_ready_case(tmp_path)
    _write_application_section(case, case_dir, heading, placeholder)
    output = tmp_path / "application.docx"

    with pytest.raises(ValueError, match=f"placeholder.*{heading}"):
        export_application(
            case_dir,
            output,
            final_approval=True,
            template_path=tmp_path / "missing-template.docx",
        )

    assert not output.exists()


def test_export_refuses_unreported_high_prior_art_risk(tmp_path: Path):
    _, case, case_dir = _write_ready_case(tmp_path)
    quality_review = case_dir / _artifact(case, "quality-review").path
    review = _official_quality_review(
        prior_art_assessment={
            "verified_disclosures": [],
            "novelty_risk": [{"claim": "C1", "level": "high"}],
            "inventive_step_risk": [],
        }
    )
    quality_review.write_text(json.dumps(review), encoding="utf-8")
    output = tmp_path / "application.docx"

    with pytest.raises(ValueError, match="official completed output recipe"):
        export_application(
            case_dir,
            output,
            final_approval=True,
            template_path=tmp_path / "missing-template.docx",
        )

    assert not output.exists()


def test_export_refuses_legacy_issues_report_even_when_issue_is_high(
    tmp_path: Path,
):
    _, case, case_dir = _write_ready_case(tmp_path)
    quality_review = case_dir / _artifact(case, "quality-review").path
    quality_review.write_text(
        json.dumps(
            {
                "issues": [
                    {
                        "issue_id": "artifact-support-gap",
                        "severity": "high",
                        "message": "claim lacks support",
                        "closed": False,
                    }
                ]
            }
        ),
        encoding="utf-8",
    )

    with pytest.raises(ValueError, match="official completed output recipe"):
        export_application(
            case_dir,
            tmp_path / "application.docx",
            final_approval=True,
            template_path=tmp_path / "missing-template.docx",
        )


def test_export_accepts_existing_scoped_final_delivery_approval_record(
    tmp_path: Path,
):
    _, case, case_dir = _write_ready_case(tmp_path)
    quality_review = case_dir / _artifact(case, "quality-review").path
    quality_review.write_text(
        json.dumps(_official_quality_review(), ensure_ascii=False),
        encoding="utf-8",
    )
    _replace_persisted_approvals(
        case_dir, [_scoped_final_delivery_approval(version=2)]
    )

    output = tmp_path / "application.docx"
    assert export_application(case_dir, output, final_approval=True) == output
    assert output.is_file()


def test_export_refuses_replaced_v2_artifacts_with_retained_v1_approval(
    tmp_path: Path,
):
    _, case, case_dir = _write_ready_case(tmp_path)
    quality_review = case_dir / _artifact(case, "quality-review").path
    quality_review.write_text(
        json.dumps(_official_quality_review(), ensure_ascii=False),
        encoding="utf-8",
    )
    _replace_persisted_approvals(
        case_dir, [_scoped_final_delivery_approval(version=1)]
    )
    output = tmp_path / "application.docx"

    with pytest.raises(ValueError, match="exact current artifact versions"):
        export_application(
            case_dir,
            output,
            final_approval=True,
            template_path=tmp_path / "missing-template.docx",
        )

    assert not output.exists()


def test_export_refuses_approval_for_a_different_application_set(tmp_path: Path):
    _, case, case_dir = _write_ready_case(tmp_path)
    quality_review = case_dir / _artifact(case, "quality-review").path
    quality_review.write_text(
        json.dumps(_official_quality_review(), ensure_ascii=False),
        encoding="utf-8",
    )
    approval = _scoped_final_delivery_approval(version=2)
    approval["application_set"] = "OTHER-CASE-v2"
    _replace_persisted_approvals(case_dir, [approval])
    output = tmp_path / "application.docx"

    with pytest.raises(ValueError, match="current application set"):
        export_application(
            case_dir,
            output,
            final_approval=True,
            template_path=tmp_path / "missing-template.docx",
        )

    assert not output.exists()


def test_workspace_and_explicit_approval_are_required_public_arguments():
    parameters = inspect.signature(export_application).parameters

    assert parameters["case_dir"].default is inspect.Parameter.empty
    assert parameters["output_path"].default is inspect.Parameter.empty
    assert parameters["final_approval"].default is inspect.Parameter.empty
    for removed_parameter in (
        "title",
        "claims",
        "sections",
        "validation_report",
        "artifacts",
    ):
        assert removed_parameter not in parameters


def test_export_rejects_missing_specification_section(tmp_path: Path):
    _, case, case_dir = _write_ready_case(tmp_path)
    specification = case_dir / _artifact(case, "specification").path
    specification.write_text(
        SPECIFICATION_TEXT.replace("\n\n附图说明\n\n" + SECTIONS["附图说明"], ""),
        encoding="utf-8",
    )

    with pytest.raises(ValueError, match="missing required section: 附图说明"):
        export_application(case_dir, tmp_path / "application.docx", final_approval=True)


def test_export_creates_output_parent_directory(tmp_path: Path):
    _, _, case_dir = _write_ready_case(tmp_path)
    output = tmp_path / "nested" / "delivery" / "application.docx"

    assert export_application(case_dir, output, final_approval=True) == output
    assert output.is_file()


def test_default_template_is_loaded_and_styles_are_applied(tmp_path: Path):
    kwargs = _export_kwargs(tmp_path)
    output = export_application(**kwargs)

    document = Document(output)
    assert TEMPLATE_PATH.is_file()
    assert "Template Sentinel" in [style.name for style in document.styles]
    paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text]
    assert paragraphs[0].text == "折叠支架"
    assert paragraphs[0].style.name == "Patent Title"
    assert next(p for p in paragraphs if p.text == "权利要求书").style.name == (
        "Patent Heading 1"
    )
    assert next(p for p in paragraphs if p.text.startswith("1.")).style.name == (
        "Patent Claim"
    )
    assert next(p for p in paragraphs if p.text == SECTIONS["技术领域"]).style.name == (
        "Patent Body"
    )
    assert next(
        p for p in paragraphs if p.text == "技术领域"
    ).paragraph_format.page_break_before
    assert next(
        p for p in paragraphs if p.text == "摘要"
    ).paragraph_format.page_break_before


def test_template_encodes_cn_patent_application_style_tokens():
    document = Document(TEMPLATE_PATH)
    section = document.sections[0]

    one_twip = 635
    assert abs(section.page_width - Cm(21.0)) <= one_twip
    assert abs(section.page_height - Cm(29.7)) <= one_twip
    assert section.top_margin == Cm(2.54)
    assert section.right_margin == Cm(2.54)
    assert section.bottom_margin == Cm(2.54)
    assert section.left_margin == Cm(2.54)

    body = document.styles["Patent Body"]
    assert "Patent Figure Caption" in [style.name for style in document.styles]
    body_fonts = body.element.rPr.rFonts
    assert body.font.name == "Times New Roman"
    assert body.font.size == Pt(11)
    assert body_fonts.get(qn("w:eastAsia")) == "SimSun"
    assert body_fonts.get(qn("w:ascii")) == "Times New Roman"
    assert body.paragraph_format.line_spacing == 1.25
    assert body.paragraph_format.space_after == Pt(6)

    title = document.styles["Patent Title"]
    heading = document.styles["Patent Heading 1"]
    assert title.font.color.rgb is not None
    assert str(title.font.color.rgb) == "000000"
    assert str(heading.font.color.rgb) == "000000"
    assert title.paragraph_format.alignment == 1
    assert heading.paragraph_format.alignment == 1

    footer_xml = section.footer._element.xml
    assert 'w:instrText xml:space="preserve">PAGE</w:instrText>' in footer_xml


def test_export_cli_reads_workspace_bound_json_input_package(tmp_path: Path):
    _, _, case_dir = _write_ready_case(tmp_path)
    package = tmp_path / "export-package.json"
    output = tmp_path / "deliverables" / "application.docx"
    package.write_text(
        json.dumps(
            {"case_dir": str(case_dir), "final_approval": True},
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    result = RUNNER.invoke(app, ["export", str(package), str(output)])

    assert result.exit_code == 0
    assert result.stdout == f"{output}\n"
    assert output.is_file()


def test_export_cli_rejects_non_boolean_final_approval(tmp_path: Path):
    _, _, case_dir = _write_ready_case(tmp_path)
    package = tmp_path / "unsafe-export-package.json"
    output = tmp_path / "application.docx"
    package.write_text(
        json.dumps({"case_dir": str(case_dir), "final_approval": "yes"}),
        encoding="utf-8",
    )

    result = RUNNER.invoke(app, ["export", str(package), str(output)])

    assert result.exit_code != 0
    assert "final_approval" in result.output
    assert not output.exists()


def test_export_cli_rejects_caller_supplied_content_and_guard_metadata(
    tmp_path: Path,
):
    _, _, case_dir = _write_ready_case(tmp_path)
    package = tmp_path / "injected-export-package.json"
    output = tmp_path / "application.docx"
    package.write_text(
        json.dumps(
            {
                "case_dir": str(case_dir),
                "final_approval": True,
                "title": "caller replacement",
                "claims": ["caller replacement"],
                "sections": SECTIONS,
                "validation_report": {"issues": []},
                "artifacts": [],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    result = RUNNER.invoke(app, ["export", str(package), str(output)])

    assert result.exit_code != 0
    assert "Extra inputs are not permitted" in result.output
    assert not output.exists()


def test_export_cli_does_not_default_missing_final_approval_to_approved(
    tmp_path: Path,
):
    _, _, case_dir = _write_ready_case(tmp_path)
    package = tmp_path / "unsafe-export-package.json"
    output = tmp_path / "application.docx"
    package.write_text(json.dumps({"case_dir": str(case_dir)}), encoding="utf-8")

    result = RUNNER.invoke(app, ["export", str(package), str(output)])

    assert result.exit_code != 0
    assert "final_approval" in result.output
    assert not output.exists()
