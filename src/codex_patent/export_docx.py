import json
import os
import re
from importlib import resources
from pathlib import Path
from typing import Literal

from docx import Document
from pydantic import BaseModel, ConfigDict, Field, ValidationError, model_validator

from codex_patent.models import ApprovalRecord, ArtifactRef, PatentCase
from codex_patent.repository import CaseRepository
from codex_patent.validation import ValidationReport


SECTION_ORDER = [
    "技术领域",
    "背景技术",
    "发明内容",
    "附图说明",
    "具体实施方式",
    "摘要",
]
SPECIFICATION_SECTION_ORDER = SECTION_ORDER[:-1]
REQUIRED_ARTIFACT_EXTENSIONS = {
    "claims": ".md",
    "specification": ".md",
    "abstract": ".md",
    "quality-review": ".json",
}
_PACKAGE_TEMPLATE = resources.files("codex_patent").joinpath(
    "templates", "cn-patent-application.docx"
)
_CHECKOUT_TEMPLATE = (
    Path(__file__).resolve().parents[2]
    / "templates"
    / "cn-patent-application.docx"
)
TEMPLATE_PATH = (
    Path(str(_PACKAGE_TEMPLATE))
    if _PACKAGE_TEMPLATE.is_file()
    else _CHECKOUT_TEMPLATE
)

QUALITY_CHECK_NAMES = (
    "claim_clarity_dependency",
    "specification_support_sufficiency",
    "cross_artifact_consistency",
    "abstract_fidelity",
    "unity",
    "subject_matter_technical_solution",
    "filing_type",
    "novelty",
    "inventive_step",
    "design_around",
    "form",
)
PLACEHOLDER_PATTERNS = (
    re.compile(r"\b(?:todo|tbd|placeholder)\b", re.IGNORECASE),
    re.compile(r"待(?:补充|确认|完善|定)"),
    re.compile(r"(?:\.{3,}|…+)"),
)


class _StrictReviewModel(BaseModel):
    model_config = ConfigDict(extra="forbid", strict=True)


class _QualityReviewInputVersions(_StrictReviewModel):
    claims: str = Field(pattern=r"^v[1-9]\d*$")
    claim_feature_map: str = Field(pattern=r"^v[1-9]\d*$")
    specification: str = Field(pattern=r"^v[1-9]\d*$")
    abstract: str = Field(pattern=r"^v[1-9]\d*$")
    drawing_plan: str = Field(pattern=r"^v[1-9]\d*$")
    prior_art: str = Field(pattern=r"^v[1-9]\d*$")


class _SourceAnchor(_StrictReviewModel):
    artifact: str | None = None
    anchor: str | None = None
    source_id: str | None = None
    locator: str | None = None
    quote: str | None = None
    document_id: str | None = None
    disclosure_anchor: str | None = None
    evidence_gap: str | None = None
    overlap: list[str] | None = None

    @model_validator(mode="after")
    def require_meaningful_anchor(self):
        if self.overlap is not None and (
            not self.overlap or any(not value.strip() for value in self.overlap)
        ):
            raise ValueError("source anchor overlap entries must be nonblank")
        text_values = (
            self.artifact,
            self.anchor,
            self.source_id,
            self.locator,
            self.quote,
            self.document_id,
            self.disclosure_anchor,
            self.evidence_gap,
        )
        if not any((value or "").strip() for value in text_values) and not self.overlap:
            raise ValueError("source anchor must identify evidence or an evidence gap")
        return self


SourceAnchor = str | _SourceAnchor


def _require_meaningful_anchors(anchors: list[SourceAnchor]) -> None:
    if any(isinstance(anchor, str) and not anchor.strip() for anchor in anchors):
        raise ValueError("source anchors must not contain blank strings")


class _QualityReviewCheck(_StrictReviewModel):
    status: Literal["completed", "not-assessable", "blocked"]
    conclusion_or_gap: str = Field(min_length=1)
    source_anchors: list[SourceAnchor]

    @model_validator(mode="after")
    def require_completed_check_evidence(self):
        _require_meaningful_anchors(self.source_anchors)
        if self.status == "completed" and not self.source_anchors:
            raise ValueError("completed checks require source anchors")
        return self


class _QualityReviewChecks(_StrictReviewModel):
    claim_clarity_dependency: _QualityReviewCheck
    specification_support_sufficiency: _QualityReviewCheck
    cross_artifact_consistency: _QualityReviewCheck
    abstract_fidelity: _QualityReviewCheck
    unity: _QualityReviewCheck
    subject_matter_technical_solution: _QualityReviewCheck
    filing_type: _QualityReviewCheck
    novelty: _QualityReviewCheck
    inventive_step: _QualityReviewCheck
    design_around: _QualityReviewCheck
    form: _QualityReviewCheck


class _QualityReviewFinding(_StrictReviewModel):
    issue_id: str = Field(min_length=1)
    severity: Literal["critical", "high", "medium", "low"]
    category: str = Field(min_length=1)
    artifact: str | None = None
    claim: str | int | None = None
    section: str | None = None
    occurrence_id: str | None = None
    location: str | None = None
    source_anchors: list[SourceAnchor]
    explanation: str = Field(min_length=1)
    suggested_action: str = Field(min_length=1)
    blocks_delivery: bool
    evidence_gap: str | None = None

    @model_validator(mode="after")
    def require_location_and_evidence(self):
        _require_meaningful_anchors(self.source_anchors)
        if not any(
            value is not None
            for value in (
                self.artifact,
                self.claim,
                self.section,
                self.occurrence_id,
                self.location,
            )
        ):
            raise ValueError("findings require an artifact, claim, section, or location")
        if not self.source_anchors and not (self.evidence_gap or "").strip():
            raise ValueError("findings require source anchors or an evidence gap")
        return self


class _OpenIssueCounts(_StrictReviewModel):
    critical: int = Field(ge=0)
    high: int = Field(ge=0)
    medium: int = Field(ge=0)
    low: int = Field(ge=0)
    total: int | None = Field(default=None, ge=0)

    @model_validator(mode="after")
    def validate_total(self):
        observed = self.critical + self.high + self.medium + self.low
        if self.total is not None and self.total != observed:
            raise ValueError("open issue total does not match severity counts")
        return self


class _InputIntegrity(_StrictReviewModel):
    status: Literal["valid-for-review"]
    all_six_readable: Literal[True]
    freshness: Literal["current"]
    mutually_version_matched: Literal[True]
    internal_identity: Literal["identifiable"]
    placeholder_block: Literal[False]


class _PriorArtRisk(_StrictReviewModel):
    claim: str | int
    level: Literal["none", "low", "medium", "high", "not-assessable"]
    basis: str | None = None
    gap: str | None = None


class _VerifiedDisclosure(_StrictReviewModel):
    document_id: str
    disclosure_anchor: str
    disclosure: str | None = None
    overlapping_features: list[str] | None = None

    @model_validator(mode="after")
    def require_exact_disclosure_identity(self):
        if not self.document_id.strip() or not self.disclosure_anchor.strip():
            raise ValueError(
                "verified disclosures require a document ID and exact disclosure anchor"
            )
        if self.disclosure is not None and not self.disclosure.strip():
            raise ValueError("verified disclosure text must not be blank")
        if self.overlapping_features is not None and (
            not self.overlapping_features
            or any(not feature.strip() for feature in self.overlapping_features)
        ):
            raise ValueError("overlapping feature identifiers must be nonblank")
        return self


class _PriorArtAssessment(_StrictReviewModel):
    verified_disclosures: list[_VerifiedDisclosure]
    novelty_risk: list[_PriorArtRisk]
    inventive_step_risk: list[_PriorArtRisk]


class _CompletedQualityReview(_StrictReviewModel):
    review_status: Literal["completed", "completed-with-issues"]
    input_versions: _QualityReviewInputVersions
    checks: _QualityReviewChecks
    findings: list[_QualityReviewFinding]
    open_issue_counts: _OpenIssueCounts
    delivery_recommendation: Literal["blocked", "ready-for-human-review"]
    unresolved_questions: list[str | dict[str, object]]
    source_anchors: list[SourceAnchor]
    application_set: str | None = None
    input_integrity: _InputIntegrity | None = None
    prior_art_assessment: _PriorArtAssessment

    @model_validator(mode="after")
    def validate_internal_gate_consistency(self):
        _require_meaningful_anchors(self.source_anchors)
        observed_counts = {
            severity: sum(
                finding.severity == severity for finding in self.findings
            )
            for severity in ("critical", "high", "medium", "low")
        }
        persisted_counts = self.open_issue_counts.model_dump(exclude={"total"})
        if persisted_counts != observed_counts:
            raise ValueError("open issue counts do not match findings")
        if self.review_status == "completed" and self.findings:
            raise ValueError("completed review cannot contain open findings")
        if self.review_status == "completed-with-issues" and not self.findings:
            raise ValueError("completed-with-issues review requires findings")
        if not self.source_anchors:
            raise ValueError("completed review requires source anchors")
        if self.prior_art_assessment is not None:
            assessment = self.prior_art_assessment
            if not assessment.verified_disclosures:
                raise ValueError(
                    "completed prior-art assessment requires verified disclosures"
                )
            verified_anchors = {
                (
                    disclosure.document_id.strip().casefold(),
                    disclosure.disclosure_anchor.strip(),
                )
                for disclosure in assessment.verified_disclosures
            }
            for check_name, risk_name in (
                ("novelty", "novelty_risk"),
                ("inventive_step", "inventive_step_risk"),
            ):
                check = getattr(self.checks, check_name)
                risks = getattr(assessment, risk_name)
                if not risks:
                    raise ValueError(
                        f"{check_name} assessment requires a claim-level risk result"
                    )
                if any(risk.level == "not-assessable" for risk in risks):
                    if check.status == "completed":
                        raise ValueError(
                            f"{check_name} cannot be completed when a risk is not-assessable"
                        )
                    continue
                if check.status == "completed":
                    check_anchors = {
                        (
                            anchor.document_id.strip().casefold(),
                            anchor.disclosure_anchor.strip(),
                        )
                        for anchor in check.source_anchors
                        if isinstance(anchor, _SourceAnchor)
                        and (anchor.document_id or "").strip()
                        and (anchor.disclosure_anchor or "").strip()
                    }
                    if not check_anchors.intersection(verified_anchors):
                        raise ValueError(
                            f"completed {check_name} check requires verified prior-art evidence"
                        )
            unreported_high_risk = any(
                risk.level == "high"
                for risk in (
                    assessment.novelty_risk
                    + assessment.inventive_step_risk
                )
            ) and not any(
                finding.severity in {"critical", "high"}
                and finding.category in {"novelty-risk", "inventive-step-risk"}
                for finding in self.findings
            )
            if unreported_high_risk:
                raise ValueError("high prior-art risk requires a blocking finding")
        return self


def _validate_content(title: str, claims: list[str], sections: dict[str, str]) -> None:
    if not title.strip():
        raise ValueError("application title must not be blank")
    if not claims or any(not claim.strip() for claim in claims):
        raise ValueError("at least one claim is required and claims must not be blank")
    for claim in claims:
        if _contains_placeholder(claim):
            raise ValueError("placeholder claim text blocks export")
    for heading in SECTION_ORDER:
        if heading not in sections:
            raise ValueError(f"missing required section: {heading}")
        if not sections[heading].strip():
            raise ValueError(f"required section must not be blank: {heading}")
        if _contains_placeholder(sections[heading]):
            raise ValueError(
                f"placeholder required section text blocks export: {heading}"
            )


def _contains_placeholder(text: str) -> bool:
    value = text.strip()
    if not any(character.isalnum() for character in value):
        return True
    return any(pattern.search(value) for pattern in PLACEHOLDER_PATTERNS)


def _paths_alias(left: Path, right: Path) -> bool:
    try:
        left_resolved = left.resolve(strict=False)
        right_resolved = right.resolve(strict=False)
    except OSError as exc:
        raise ValueError("unable to validate export output path") from exc
    if os.path.normcase(str(left_resolved)) == os.path.normcase(str(right_resolved)):
        return True
    try:
        return left.exists() and right.exists() and os.path.samefile(left, right)
    except OSError:
        return False


def _validate_output_path(
    case_dir: Path,
    case: PatentCase,
    output_path: Path,
    template_path: Path | None,
) -> Path:
    output_path = Path(output_path)
    protected_paths = [path for path in case_dir.rglob("*") if path.is_file()]
    protected_paths.extend(case_dir / artifact.path for artifact in case.artifacts)
    protected_paths.append(Path(template_path) if template_path else TEMPLATE_PATH)
    if any(_paths_alias(output_path, protected) for protected in protected_paths):
        raise ValueError("output path collides with a protected case or template path")
    return output_path


def _load_case(case_dir: Path) -> PatentCase:
    case_dir = Path(case_dir)
    if not case_dir.is_dir():
        raise ValueError(f"case workspace does not exist: {case_dir}")
    try:
        case = CaseRepository(case_dir.parent).load(case_dir.name)
    except (
        FileNotFoundError,
        OSError,
        json.JSONDecodeError,
        ValidationError,
    ) as exc:
        raise ValueError(f"invalid case workspace: {case_dir}") from exc
    if case.case_id != case_dir.name:
        raise ValueError("persisted case ID does not match the case directory")
    return case


def _current_required_artifacts(case: PatentCase) -> dict[str, ArtifactRef]:
    selected: dict[str, ArtifactRef] = {}
    for artifact_type in REQUIRED_ARTIFACT_EXTENSIONS:
        candidates = [
            artifact
            for artifact in case.artifacts
            if artifact.artifact_type.casefold() == artifact_type
        ]
        if not candidates:
            raise ValueError(f"missing required artifact: {artifact_type}")
        current = [artifact for artifact in candidates if not artifact.stale]
        if not current:
            raise ValueError(f"stale {artifact_type} artifact blocks export")
        if len(current) > 1:
            raise ValueError(f"multiple current {artifact_type} artifacts block export")
        selected[artifact_type] = current[0]

    versions = {artifact.version for artifact in selected.values()}
    if len(versions) != 1:
        raise ValueError("required artifacts must share one version")
    return selected


def _read_artifact(
    case_dir: Path,
    artifact_type: str,
    artifact: ArtifactRef,
) -> str:
    relative_path = Path(artifact.path)
    expected_name = (
        f"{artifact_type}-v{artifact.version}"
        f"{REQUIRED_ARTIFACT_EXTENSIONS[artifact_type]}"
    )
    if relative_path.name != expected_name:
        raise ValueError(
            f"{artifact_type} artifact path must name {expected_name}"
        )
    if relative_path.is_absolute():
        raise ValueError(
            f"{artifact_type} artifact must resolve inside the case workspace"
        )

    resolved_case_dir = case_dir.resolve()
    artifact_path = (case_dir / relative_path).resolve()
    if not artifact_path.is_relative_to(resolved_case_dir):
        raise ValueError(
            f"{artifact_type} artifact must resolve inside the case workspace"
        )
    if not artifact_path.is_file():
        raise ValueError(f"missing {artifact_type} artifact file: {artifact.path}")
    try:
        content = artifact_path.read_text(encoding="utf-8")
    except UnicodeDecodeError as exc:
        raise ValueError(
            f"unreadable {artifact_type} artifact: {artifact.path}"
        ) from exc
    if not content.strip():
        raise ValueError(f"empty {artifact_type} artifact blocks export")
    return content


def _parse_claims(content: str) -> list[str]:
    claims = [
        block.strip()
        for block in re.split(r"\n\s*\n", content)
        if block.strip()
    ]
    if not claims:
        raise ValueError("at least one claim is required")
    return claims


def _parse_specification(content: str) -> dict[str, str]:
    lines = content.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    positions: dict[str, int] = {}
    for index, line in enumerate(lines):
        heading = line.strip()
        if heading not in SPECIFICATION_SECTION_ORDER:
            continue
        if heading in positions:
            raise ValueError(f"duplicate required section: {heading}")
        positions[heading] = index

    for heading in SPECIFICATION_SECTION_ORDER:
        if heading not in positions:
            raise ValueError(f"missing required section: {heading}")
    observed_positions = [
        positions[heading] for heading in SPECIFICATION_SECTION_ORDER
    ]
    if observed_positions != sorted(observed_positions):
        raise ValueError("specification sections are not in the required order")

    sections: dict[str, str] = {}
    for index, heading in enumerate(SPECIFICATION_SECTION_ORDER):
        start = positions[heading] + 1
        end = (
            positions[SPECIFICATION_SECTION_ORDER[index + 1]]
            if index + 1 < len(SPECIFICATION_SECTION_ORDER)
            else len(lines)
        )
        sections[heading] = "\n".join(lines[start:end]).strip()
    return sections


def _quality_review_report(content: str, version: int) -> _CompletedQualityReview:
    try:
        value = json.loads(content)
    except json.JSONDecodeError as exc:
        raise ValueError(
            "quality-review artifact must contain a valid JSON object"
        ) from exc
    if not isinstance(value, dict):
        raise ValueError("quality-review artifact must contain a valid JSON object")
    try:
        review = _CompletedQualityReview.model_validate(value)
    except ValidationError as exc:
        raise ValueError(
            "quality-review artifact must follow the official completed output recipe"
        ) from exc

    expected_version = f"v{version}"
    if any(
        input_version != expected_version
        for input_version in review.input_versions.model_dump().values()
    ):
        raise ValueError(
            "quality-review artifact does not cover the exact current artifact versions"
        )
    if any(
        getattr(review.checks, check_name).status != "completed"
        for check_name in QUALITY_CHECK_NAMES
    ):
        raise ValueError("quality-review artifact is not eligible for export")
    if (
        review.delivery_recommendation != "ready-for-human-review"
        or review.open_issue_counts.critical != 0
        or review.open_issue_counts.high != 0
        or any(finding.blocks_delivery for finding in review.findings)
    ):
        raise ValueError("quality-review artifact is not eligible for export")
    return review


def _scoped_final_delivery_approval(
    case: PatentCase,
    artifacts: dict[str, ArtifactRef],
) -> ApprovalRecord:
    approvals = [
        approval
        for approval in case.approvals
        if isinstance(approval, ApprovalRecord)
        and approval.type == "final-delivery"
        and approval.status == "approved"
        and approval.current is True
    ]
    if not approvals:
        raise ValueError("persisted final-delivery approval must be scoped")
    if len(approvals) != 1:
        raise ValueError("multiple current final-delivery approvals block export")

    approval = approvals[0]
    expected_scope = {
        "claims": f"v{artifacts['claims'].version}",
        "specification": f"v{artifacts['specification'].version}",
        "abstract": f"v{artifacts['abstract'].version}",
        "quality_review": f"v{artifacts['quality-review'].version}",
        "action": "DOCX export",
    }
    if approval.scope.model_dump() != expected_scope:
        raise ValueError(
            "final-delivery approval must cover the exact current artifact versions "
            "and DOCX export action"
        )
    current_version = artifacts["claims"].version
    expected_application_set = f"{case.case_id}-v{current_version}"
    if approval.application_set != expected_application_set:
        raise ValueError(
            "final-delivery approval must cover the current application set"
        )
    return approval


def _open_template(template_path: Path | None) -> Document:
    if template_path is not None:
        return Document(template_path)
    if _PACKAGE_TEMPLATE.is_file():
        with resources.as_file(_PACKAGE_TEMPLATE) as resolved_template:
            return Document(resolved_template)
    return Document(_CHECKOUT_TEMPLATE)


def export_application(
    case_dir: Path,
    output_path: Path,
    *,
    final_approval: bool,
    template_path: Path | None = None,
) -> Path:
    if type(final_approval) is not bool:
        raise ValueError("explicit final approval must be an actual boolean")
    if final_approval is not True:
        raise ValueError("explicit final approval is required")

    case_dir = Path(case_dir)
    case = _load_case(case_dir)
    output_path = _validate_output_path(
        case_dir,
        case,
        Path(output_path),
        template_path,
    )
    validation_report = ValidationReport(issues=case.issues)
    if not validation_report.eligible_for_final_delivery:
        raise ValueError("open high-severity validation issues block export")
    if any(
        artifact.artifact_type.casefold() == "docx" and artifact.stale
        for artifact in case.artifacts
    ):
        raise ValueError("stale docx artifact blocks export")

    artifacts = _current_required_artifacts(case)
    approval = _scoped_final_delivery_approval(case, artifacts)
    contents = {
        artifact_type: _read_artifact(case_dir, artifact_type, artifact)
        for artifact_type, artifact in artifacts.items()
    }
    review = _quality_review_report(
        contents["quality-review"], artifacts["quality-review"].version
    )
    if (
        review.application_set is not None
        and review.application_set != approval.application_set
    ):
        raise ValueError(
            "final-delivery approval does not cover the reviewed application set"
        )
    claims = _parse_claims(contents["claims"])
    sections = _parse_specification(contents["specification"])
    sections["摘要"] = contents["abstract"].strip()
    _validate_content(case.title, claims, sections)

    document = _open_template(template_path)
    document.add_paragraph(case.title.strip(), style="Patent Title")
    document.add_paragraph("权利要求书", style="Patent Heading 1")
    for claim in claims:
        document.add_paragraph(claim, style="Patent Claim")
    for heading in SECTION_ORDER:
        heading_paragraph = document.add_paragraph(heading, style="Patent Heading 1")
        if heading in {"技术领域", "摘要"}:
            heading_paragraph.paragraph_format.page_break_before = True
        for block in re.split(r"\n\s*\n", sections[heading]):
            if block.strip():
                document.add_paragraph(block.strip(), style="Patent Body")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path
